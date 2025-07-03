import streamlit as st
from datetime import datetime
import os
from gemini_client import GeminiClient
import PyPDF2
import docx

# 페이지 설정
st.set_page_config(
    page_title="서던포스트 AI",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# CSS 스타일링
st.markdown("""
<style>
    /* 전체 배경 */
    .main {
        background-color: #f7f7f8;
    }
    
    /* 헤더 스타일 */
    .header {
        background: linear-gradient(90deg, #1f77b4, #ff7f0e);
        padding: 1rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        text-align: center;
        color: white;
    }
    
    /* 채팅 컨테이너 */
    .chat-container {
        background: white;
        border-radius: 10px;
        padding: 1rem;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        margin-bottom: 1rem;
    }
    
    /* 파일 업로드 영역 */
    .upload-area {
        background: #f8f9fa;
        border: 2px dashed #dee2e6;
        border-radius: 10px;
        padding: 2rem;
        text-align: center;
        margin-bottom: 2rem;
    }
    
    /* 사이드바 스타일 */
    .css-1d391kg {
        background-color: #f8f9fa;
    }
    
    /* 버튼 스타일 */
    .stButton > button {
        background: linear-gradient(90deg, #1f77b4, #ff7f0e);
        color: white;
        border: none;
        border-radius: 20px;
        padding: 0.5rem 1rem;
    }
    
    /* 입력창 스타일 */
    .stTextInput > div > div > input {
        border-radius: 20px;
        border: 2px solid #dee2e6;
    }
    
    /* 채팅 메시지 스타일 */
    .user-message {
        background: #1f77b4;
        color: white;
        padding: 1rem;
        border-radius: 15px 15px 0 15px;
        margin: 0.5rem 0;
        max-width: 80%;
        margin-left: auto;
    }
    
    .assistant-message {
        background: #f8f9fa;
        color: #333;
        padding: 1rem;
        border-radius: 15px 15px 15px 0;
        margin: 0.5rem 0;
        max-width: 80%;
        border: 1px solid #dee2e6;
    }
    
    /* 파일 미리보기 */
    .file-preview {
        background: #e9ecef;
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
        border-left: 4px solid #1f77b4;
    }
    
    /* 로딩 스피너 */
    .loading {
        text-align: center;
        padding: 2rem;
        color: #666;
    }
</style>
""", unsafe_allow_html=True)

# 세션 상태 초기화
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'gemini_client' not in st.session_state:
    st.session_state.gemini_client = None
if 'uploaded_file_text' not in st.session_state:
    st.session_state.uploaded_file_text = ""
if 'uploaded_file_name' not in st.session_state:
    st.session_state.uploaded_file_name = None

# 파일 텍스트 추출 함수
@st.cache_data(show_spinner=False)
def extract_text_from_file(uploaded_file):
    if uploaded_file is None:
        return ""
    try:
        filetype = uploaded_file.name.split('.')[-1].lower()
        if filetype == 'pdf':
            reader = PyPDF2.PdfReader(uploaded_file)
            texts = [page.extract_text() or '' for page in reader.pages]
            non_empty_texts = [t for t in texts if t and t.strip()]
            text = "\n".join(non_empty_texts)
            return text
        elif filetype in ['docx', 'doc']:
            doc = docx.Document(uploaded_file)
            texts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            # 표의 셀 텍스트도 추출
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            texts.append(cell_text)
            text = "\n".join(texts)
            print("[docx 추출] 전체 길이:", len(text))
            print("[docx 추출] 앞 500자:\n", text[:500])
            print("[docx 추출] 끝 500자:\n", text[-500:])
            return text
        elif filetype in ['txt', 'md']:
            return uploaded_file.read().decode('utf-8')
        else:
            return "[지원하지 않는 파일 형식입니다: PDF, DOCX, TXT만 지원]"
    except Exception as e:
        import traceback
        print("파일 읽기 오류:", traceback.format_exc())
        return f"[파일 읽기 오류: {str(e)}]"

# 젠스파크 명령어 감지
GENSPARK_TRIGGER = ["젠스파크", "genspark", "슬라이드", "목차", "프롬프트"]
def is_genspark_command(user_input):
    return any(trigger in user_input.lower() for trigger in GENSPARK_TRIGGER)

def make_genspark_prompt(chat_history, file_text):
    chat_summary = "\n".join([f"[사용자] {m['content']}" for m in chat_history if m['role']=='user'])
    prompt = f"""
아래 [제안요청서 원문]의 내용을 반드시 근거로 하여, 젠스파크(Genspark) 프레젠테이션용 슬라이드 목차와 각 슬라이드별 AI 프롬프트를 매우 고급스럽고 디테일하게 작성해 주세요.
- 임의로 가정하거나 예시를 들지 말고, 반드시 실제 원문 내용을 반영하세요.
- 회사명은 '서던포스트'로 고정하세요.

[제안요청서 원문]
{file_text[:50000]}

[대화에서 얻은 추가 정보]
{chat_summary}

요구사항:
1. 논리적 구조, 설득력, 차별화 포인트, 실무적 디테일을 모두 포함
2. 각 슬라이드별로 제목, 목적, 주요 메시지, 구체적 프롬프트(젠스파크 AI에 입력할 수준)를 명확하게 구분
3. 최신 입찰/제안 트렌드와 경쟁사 대비 강점, 실질적 실행방안, 수치/사례 등도 적극 반영
4. 전체 목차는 8~15개 슬라이드로 구성하며, 각 항목은 한글로 작성

형식:
---
## 젠스파크용 슬라이드 목차 및 프롬프트 (서던포스트)

### 1. [슬라이드 제목]
- 목적: [이 슬라이드의 핵심 목적]
- 주요 메시지: [핵심 설득 포인트]
- 프롬프트: [젠스파크 AI에 입력할 상세 프롬프트]

### 2. [슬라이드 제목]
- 목적: ...
- 주요 메시지: ...
- 프롬프트: ...

...
---
"""
    print("==== AI에 전달되는 최종 프롬프트(앞부분) ====")
    print(prompt[:1000])
    return prompt

# 메인 UI
def main():
    # 헤더
    st.markdown("""
    <div class="header">
        <h1>🤖 서던포스트 AI</h1>
        <p>입찰제안서 분석 • 젠스파크 슬라이드 생성 • 문서 작성</p>
    </div>
    """, unsafe_allow_html=True)

    # 파일 업로드 (최상위, columns 밖!)
    uploaded_file = st.file_uploader(
        "제안요청서 업로드 (PDF, DOCX, TXT, MD)",
        type=["pdf", "docx", "doc", "txt", "md"],
        help="PDF, DOCX, TXT, MD 파일을 업로드하세요"
    )
    if uploaded_file:
        st.session_state.uploaded_file_text = extract_text_from_file(uploaded_file)
        st.session_state.uploaded_file_name = uploaded_file.name
        st.success(f"✅ {uploaded_file.name}")
        with st.expander("📄 파일 미리보기"):
            file_text = st.session_state.uploaded_file_text
            # 앞부분이 비어있고 뒷부분에만 내용이 있으면 안내
            if file_text and len(file_text) > 100 and file_text[:100].strip() == '':
                st.info("⚠️ 파일 앞부분이 비어 있습니다. 실제 내용은 아래에서부터 시작됩니다.")
            if (
                not file_text or
                not file_text.strip() or
                file_text.strip().startswith("[파일 읽기 오류:") or
                file_text.strip().startswith("[지원하지 않는 파일")
            ):
                st.error("❌ 파일에서 텍스트를 추출하지 못했습니다. PDF/DOCX/TXT/MD 파일이 맞는지, 암호화/스캔본이 아닌지 확인하세요!\n\n" + (file_text if file_text else ""))
            else:
                st.text_area(
                    "파일 내용",
                    file_text[:50000] + "..." if len(file_text) > 50000 else file_text,
                    height=500,
                    disabled=True
                )

    # 사이드바 (API 키 설정, 대화 이력 다운로드)
    with st.sidebar:
        st.title("⚙️ 설정")
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            api_key = st.text_input(
                "Gemini API 키",
                value="AIzaSyBQTovczMgxdVExbcstS7A-aLB2regABM4",
                type="password",
                help="Google AI Studio에서 발급받은 API 키를 입력하세요."
            )
        if api_key and not st.session_state.gemini_client:
            try:
                st.session_state.gemini_client = GeminiClient(api_key)
                st.success("✅ API 연결 완료")
            except Exception as e:
                st.error(f"❌ API 오류: {str(e)}")
        st.markdown("---")
        if st.session_state.chat_history:
            st.subheader("💾 내보내기")
            chat_text = "\n\n".join([
                f"[{m['role'].upper()}] {m['content']}" for m in st.session_state.chat_history
            ])
            st.download_button(
                label="📥 대화 내역 다운로드",
                data=chat_text,
                file_name=f"서던포스트_AI_대화_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain"
            )

    # 채팅 메시지 표시 (columns 사용 가능, 입력 위젯 X)
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        for msg in st.session_state.chat_history:
            if msg["role"] == "user":
                st.markdown(f"""
                <div class="user-message">
                    {msg["content"]}
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div class="assistant-message">
                    {msg["content"]}
                </div>
                """, unsafe_allow_html=True)

    # chat_input은 main의 마지막에 단독으로 위치!
    if st.session_state.gemini_client:
        user_input = st.chat_input(
            placeholder="예: 젠스파크 슬라이드 목차 생성해줘, 이 문서 분석해줘"
        )
        if user_input:
            st.session_state.chat_history.append({"role": "user", "content": user_input})
            # 젠스파크 명령 처리
            if is_genspark_command(user_input):
                file_text = st.session_state.uploaded_file_text
                if not file_text or not file_text.strip() or len(file_text.strip()) < 10:
                    st.error("⚠️ 제안요청서 파일이 정상적으로 읽히지 않았습니다. PDF/DOCX/TXT/MD 파일이 맞는지, 내용이 충분한지 확인해주세요!")
                else:
                    # 실제로 AI에 전달되는 텍스트 일부를 로그로 출력
                    print("==== AI 전달 텍스트 길이:", len(file_text))
                    print("==== AI 전달 텍스트 앞 500자 ====")
                    print(file_text[:500])
                    print("==== AI 전달 텍스트 끝 500자 ====")
                    print(file_text[-500:])
                    with st.spinner("🤖 젠스파크 슬라이드 목차/프롬프트 생성 중..."):
                        prompt = make_genspark_prompt(st.session_state.chat_history, file_text)
                        result = st.session_state.gemini_client.generate_chat([
                            {"role": "user", "content": prompt}
                        ], temperature=0.5)
                        print("==== Gemini API 응답 ====")
                        print(result)
                        if result["success"]:
                            st.session_state.chat_history.append({"role": "assistant", "content": result["content"]})
                            # 이어 생성 버튼 표시
                            if result.get("reached_limit"):
                                if st.button("이어 생성"):
                                    with st.spinner("🤖 이어서 생성 중..."):
                                        more_result = st.session_state.gemini_client.generate_chat(
                                            st.session_state.chat_history, temperature=0.5
                                        )
                                        if more_result["success"]:
                                            st.session_state.chat_history.append({"role": "assistant", "content": more_result["content"]})
                                            st.rerun()
                                        else:
                                            st.error(f"❌ 이어 생성 오류: {more_result['error']}")
                            st.rerun()
                        else:
                            st.error(f"❌ 오류: {result['error']}")
            else:
                # 일반 채팅
                with st.spinner("🤖 AI가 답변을 생성 중입니다..."):
                    result = st.session_state.gemini_client.generate_chat(st.session_state.chat_history)
                    if result["success"]:
                        st.session_state.chat_history.append({"role": "assistant", "content": result["content"]})
                        st.rerun()
                    else:
                        st.error(f"❌ 오류: {result['error']}")
    else:
        st.info("💡 사이드바에서 API 키를 설정해주세요.")
        st.markdown("""
        <div class="assistant-message">
            안녕하세요! 서던포스트 AI입니다. 👋<br>
            <b>사용 방법:</b><br>
            1. 사이드바에서 API 키를 설정하세요<br>
            2. 제안요청서 파일을 업로드하세요<br>
            3. 채팅으로 요청하세요!<br>
            <b>주요 기능:</b><br>
            • 젠스파크 슬라이드 목차/프롬프트 생성<br>
            • 입찰제안서 분석 및 작성<br>
            • 문서 개선 및 수정<br>
            <br>모든 답변에 회사명 '서던포스트'가 자동으로 반영됩니다.
        </div>
        """, unsafe_allow_html=True)

if __name__ == "__main__":
    main() 
